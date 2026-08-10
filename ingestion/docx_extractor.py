from pathlib import Path

from docx import Document


class DocxExtractor:

    def extract(self, docx_file):

        docx_file = Path(docx_file)

        if not docx_file.exists():
            raise FileNotFoundError(
                f"DOCX not found: {docx_file}"
            )

        document = Document(
            docx_file
        )

        parts = []

        # Paragraphs

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # ----------------------------------
        # Tables
        # ----------------------------------

        for table in document.tables:

            parts.append("[TABLE]")

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    text = (
                        cell.text
                        .strip()
                        .replace("\n", " ")
                    )

                    cells.append(text)

                parts.append(
                    " | ".join(cells)
                )

        full_text = "\n".join(parts)

        return {
            "text": full_text,
            "metadata": {
                "file_name": docx_file.name,
                "extraction_method": "docx",
                "paragraphs": len(
                    document.paragraphs
                ),
                "tables": len(
                    document.tables
                )
            }
        }